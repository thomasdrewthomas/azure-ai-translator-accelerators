"""
Module for reading documents from URLs.

This module provides functions to read DOCX and PDF files from given URLs and extract their text content.
"""

import logging
import io
from io import BytesIO
import requests
from docx import Document
import fitz  # PyMuPDF


def read_docx_from_url(docx_url):
    """
    Reads a DOCX file from the provided URL and extracts its text content.

    Args:
        docx_url (str): The URL of the DOCX file to read.

    Returns:
        str: The extracted text content from the DOCX file.
    """
    logging.info("Attempting to fetch DOCX from URL: %s", docx_url)
    response = requests.get(docx_url, timeout=30)
    response.raise_for_status()  # Ensure the request succeeded
    logging.info("DOCX file fetched successfully.")

    doc = Document(BytesIO(response.content))
    full_text = []

    try:
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text.append(paragraph.text)

        combined_text = "\n".join(filter(None, full_text))
        logging.info("Text extracted from DOCX file successfully.")
        return combined_text
    except Exception as e:
        logging.error("Error processing text from DOCX: %s", e, exc_info=True)
        raise


def read_pdf_from_url(pdf_url):
    """
    Reads a PDF file from the provided URL and extracts its text content.

    Args:
        pdf_url (str): The URL of the PDF file to read.

    Returns:
        str: The extracted text content from the PDF file.
    """
    logging.info("Attempting to fetch PDF from URL: %s", pdf_url)
    response = requests.get(pdf_url, timeout=30)
    response.raise_for_status()  # Ensure the request succeeded
    logging.info("PDF file fetched successfully.")

    doc = fitz.open(stream=io.BytesIO(response.content), filetype="pdf")
    text = "".join(page.get_text() for page in doc)
    logging.info("Text extracted from PDF file successfully.")
    return text
