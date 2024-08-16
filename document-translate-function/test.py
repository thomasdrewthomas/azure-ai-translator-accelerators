import os
import json
from docx import Document
import fitz  # PyMuPDF
from openai import AzureOpenAI
from azure.identity import DefaultAzureCredential, get_bearer_token_provider

API_KEY = "907e10f8f9c341a0aaaaf0ef827395e3"
AZURE_OPENAI_ENDPOINT = "https://az-demo.openai.azure.com/"
CHAT_COMPLETIONS_DEPLOYMENT_NAME = "az-chat"

def read_docx(file_path):
    """Reads a .docx file and returns its text content, including tables and text boxes."""
    doc = Document(file_path)
    full_text = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text.append(paragraph.text)

    # Check if the document has text boxes via the document's XML
    for shape in doc.element.xpath('.//w:txbxContent//w:p'):
        full_text.append(shape.text)

    return '\n'.join(filter(None, full_text))  # filter removes empty strings


def read_pdf(file_path):
    """Reads a .pdf file and returns its text content."""
    doc = fitz.open(file_path)
    return "".join(page.get_text() for page in doc)


def get_gpt_response(prompt_text, system_prompt, few_shot_examples, chat_parameters):
    """Sends a request to the Azure OpenAI API and returns the response."""
    client = AzureOpenAI(
        api_key=API_KEY,
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        api_version="2024-02-01",
    )

    messages = [{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt_text}]
    messages.extend({"role": "assistant", "content": example["chatbotResponse"]} for example in few_shot_examples)
    messages.extend({"role": "user", "content": example["userInput"]} for example in few_shot_examples)

    completion = client.chat.completions.create(
        model=chat_parameters["deploymentName"],
        messages=messages,
        max_tokens=chat_parameters.get("maxResponseLength", 800),
        temperature=chat_parameters.get("temperature", 0.7),
        top_p=chat_parameters.get("topProbabilities", 0.95),
        stop=chat_parameters.get("stopSequences"),
        frequency_penalty=chat_parameters.get("frequencyPenalty", 0),
        presence_penalty=chat_parameters.get("presencePenalty", 0)
    )

    return completion.to_json()


def parse_response(json_data):
    """Parses the JSON response from the API to extract addresses."""
    data = json.loads(json_data)
    if 'choices' in data and data['choices']:
        text = data['choices'][0]['message']['content']
        return [line.strip() for line in text.strip().split('\n') if line.strip()]
    return []


def process_file(file_path):
    """Processes a .docx or .pdf file to extract addresses using the OpenAI API."""
    if file_path.endswith('.docx'):
        text = read_docx(file_path)
    elif file_path.endswith('.pdf'):
        text = read_pdf(file_path)
    else:
        raise ValueError("Unsupported file type. Please provide a .docx or .pdf file.")
    # print(text)

    response = get_gpt_response(text, system_prompt, few_shot_examples, chat_parameters)
    return parse_response(response)


# Set up parameters and examples
system_prompt = "- Extract all location addresses from the provided text. \n- Maintain the original address format. If the address spans multiple lines, keep it multiline. \n- Do not translate or modify the content. \n- Extract each line of the address in a separate line. \n- Provide only the extracted addresses without adding any additional text.\n"
few_shot_examples = []  # Define examples if necessary
chat_parameters = {
    "deploymentName": "az-chat",
    "maxResponseLength": 800,
    "temperature": 0.7,
    "topProbabilities": 0.95,
    "stopSequences": None,
    "pastMessagesToInclude": 10,
    "frequencyPenalty": 0,
    "presencePenalty": 0
}

# Example usage for a file
file_path = "1PolishOriginal124.docx"
# file_path = "tst.pdf"
#file_path = "2-Polish-Translated.docx"

addresses = process_file(file_path)
print(addresses)